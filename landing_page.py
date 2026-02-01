import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import io
import random
from datetime import datetime

# --- 1. DATA PARSER ---
@st.cache_data
def load_custom_data():
    data_map = {
        "USERS": {}, "DESIG": [], "TRADE": [], 
        "G_TYPE": [], "AUTH_Y": [], "AUTH_Z": []
    }
    if not os.path.exists("data.txt"):
        return data_map
    try:
        with open("data.txt", "r", encoding="utf-8-sig") as f:
            current_section = None
            for line in f:
                clean_line = line.strip()
                if not clean_line: continue
                if clean_line == "USER_LIST": current_section = "USERS"
                elif clean_line == "DESIGNATIONS": current_section = "DESIG"
                elif clean_line == "TRADES": current_section = "TRADE"
                elif clean_line == "GRIEVANCE_TYPES": current_section = "G_TYPE"
                elif clean_line == "AUTHORITIES_Y": current_section = "AUTH_Y"
                elif clean_line == "AUTHORITIES_Z": current_section = "AUTH_Z"
                elif current_section == "USERS" and "," in clean_line:
                    uid, uname = clean_line.split(",", 1)
                    data_map["USERS"][uid.strip().upper()] = uname.strip()
                elif current_section:
                    data_map[current_section].append(clean_line)
    except Exception: pass
    return data_map

data = load_custom_data()

# --- 2. STYLING ---
st.set_page_config(page_title="Grievance Management System", layout="wide")
st.markdown("""
    <style>
    .stApp { background-color: #273342; color: #e2e8f0; }
    .welcome-text { font-size: 2.5rem !important; color: #3b82f6 !important; font-weight: 800; }
    label { color: #60a5fa !important; font-weight: 700 !important; font-size: 1.5rem !important; }
    .section-header { color: #ffffff; font-size: 2.2rem; font-weight: 800; border-bottom: 3px solid #3b82f6; }
    input, div[data-baseweb="select"] > div, textarea {
        background-color: #ffffff !important; color: #1e293b !important;
        border: 2px solid #3b82f6 !important; border-radius: 6px !important;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 3. LOGIN ---
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    _, col_mid, _ = st.columns([0.5, 1.2, 0.5])
    with col_mid:
        if os.path.exists("banner.png"): st.image("banner.png", use_container_width=True)
        st.markdown("<h1 style='text-align: center;'>LOGIN</h1>", unsafe_allow_html=True)
        login_id = st.text_input("Enter Login Credentials", type="password").upper().strip()
        if st.button("ENTER"):
            clean_login = re.sub(r'[^A-Z0-9]', '', login_id)
            if clean_login in data["USERS"]:
                st.session_state["authenticated"] = True
                st.session_state["user_name"] = data["USERS"][clean_login]
                st.rerun()
            else: st.error("Invalid ID")
    st.stop()

# --- 4. WORD DOCUMENT LOGIC ---
def generate_official_docx(form_data, user_name, grievance_id):
    doc = Document()
    
    # 1. Header with Logo and ID
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, Inches(6))
    htab_cells = htable.rows[0].cells
    
    # Add Logo to Header
    if os.path.exists("logo.png"):
        run = htab_cells[0].paragraphs[0].add_run()
        run.add_picture("logo.png", width=Inches(0.8))
    
    # Add Grievance ID to Header
    id_para = htab_cells[1].paragraphs[0]
    id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    id_run = id_para.add_run(f"Ref ID: {grievance_id}")
    id_run.font.bold = True
    id_run.font.size = Pt(10)

    # 2. Main Titles
    title = doc.add_heading('उत्तर रेलवे - कैरिज वर्कशॉप आलमाग', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Grievance Redressal Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Content Table (For clean alignment)
    table = doc.add_table(rows=0, cols=2)
    
    def add_row(label, value):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    add_row("Grievance दिनांक:", form_data['date'])
    add_row("कर्मचारी का नाम:", form_data['name'])
    add_row("HRMS ID:", form_data['hrms'])
    add_row("पद:", form_data['desig'])
    add_row("ट्रेड:", form_data['trade'])
    add_row("Employee No:", form_data['emp_no'])
    add_row("सेक्शन:", form_data['section'])
    add_row("Grievance प्रकार:", form_data['type'])
    add_row("संबंधित अधिकारी (To):", form_data['y'])
    add_row("जारीकर्ता अधिकारी (By):", form_data['z'])

    doc.add_paragraph("\nविवरण (Detailed Grievance):").runs[0].font.bold = True
    doc.add_paragraph(form_data['detail'])

    # 4. Footer
    doc.add_paragraph("\n" * 2)
    footer_para = doc.add_paragraph(f"Employee/ Officer registering grievance: {user_name}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.runs[0].font.bold = True

    # Save to BytesIO
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- 5. MAIN UI ---
col_logo, col_title = st.columns([0.15, 0.85])
with col_logo:
    if os.path.exists("logo.png"): st.image("logo.png", width=120)
with col_title:
    st.markdown("<h1 style='color: white;'>Carriage Workshop Alambagh, Lucknow</h1>", unsafe_allow_html=True)

st.markdown(f'<p class="welcome-text">Welcome, {st.session_state["user_name"]} 👋</p>', unsafe_allow_html=True)

with st.form("main_form"):
    st.markdown('<div class="section-header">📋 कर्मचारी का विवरण (Employee details)</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        emp_name = st.text_input("1. कर्मचारी का नाम")
        emp_desig = st.selectbox("2. कर्मचारी का पद", data["DESIG"])
        emp_trade = st.selectbox("3. कर्मचारी का ट्रेड", data["TRADE"])
    with c2:
        emp_no = st.text_input("4. Employee Number")
        hrm_id = st.text_input("5. HRMS ID", max_chars=6).upper()
        section = st.text_input("6. सेक्शन")

    st.markdown('<div class="section-header">📝 समस्या विवरण (Grievance)</div>', unsafe_allow_html=True)
    gx, gy = st.columns(2)
    with gx:
        g_type = st.selectbox("Grievance प्रकार", data["G_TYPE"])
        auth_y = st.selectbox("संबंधित अधिकारी (Letter To)", data["AUTH_Y"])
    with gy:
        date_c = st.date_input("Grievance दिनांक")
        auth_z = st.selectbox("पत्र जारीकर्ता अधिकारी (Letter By)", data["AUTH_Z"])
    
    g_detail = st.text_area("विवरण (Detailed Grievance)")
    
    _, btn_col, _ = st.columns([1, 1, 1])
    with btn_col:
        if os.path.exists("button.png"): st.image("button.png", use_container_width=True)
        submit = st.form_submit_button("GENERATE LETTER")

if submit:
    if not emp_name or not hrm_id:
        st.error("Fill Name and HRMS ID")
    else:
        g_id = f"CWM/Grievance/{emp_desig}/{hrm_id}/"
        pdf_data = {
            "date": date_c.strftime("%d-%m-%Y"),
            "name": emp_name, "desig": emp_desig, "trade": emp_trade,
            "emp_no": emp_no, "hrms": hrm_id, "section": section,
            "type": g_type, "detail": g_detail, "y": auth_y, "z": auth_z
        }
        try:
            doc_bytes = generate_official_docx(pdf_data, st.session_state["user_name"], g_id)
            st.success(f"✅ Generated ID: {g_id}")
            file_name_clean = g_id.replace("/", "_")
            st.download_button("📥 Download Word File", doc_bytes, f"{file_name_clean}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"Error: {e}")

